import hashlib
import json
import os
import shutil
import subprocess
import sys
import tempfile
from dataclasses import dataclass
from pathlib import Path
from typing import Callable, Iterable, List

from app.ingest.models import SourceDocument

TEXT_LIKE_SUFFIXES = {".txt", ".md", ".csv", ".log"}
COMMON_TEXT_ENCODINGS = ("utf-8", "gb18030", "gbk")
WINDOWS_COMMON_SOFFICE_PATHS = (
    r"C:\Program Files\LibreOffice\program\soffice.exe",
    r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
)
WINDOWS_COMMON_ANTIWORD_PATHS = (
    r"C:\Program Files\antiword\antiword.exe",
    r"C:\Program Files (x86)\antiword\antiword.exe",
)
WINDOWS_COMMON_CATDOC_PATHS = (
    r"C:\Program Files\catdoc\catdoc.exe",
    r"C:\Program Files (x86)\catdoc\catdoc.exe",
)


@dataclass
class PDFParseOptions:
    ocr_fallback_enabled: bool = True
    text_min_chars: int = 30
    ocr_engine: str = "tesseract"
    ocr_lang: str = "chi_sim+eng"
    ocr_dpi: int = 200
    ocr_max_pages: int = 0
    ocr_tesseract_cmd: str = ""


def _decode_text_bytes(raw: bytes) -> str:
    for encoding in COMMON_TEXT_ENCODINGS:
        try:
            return raw.decode(encoding=encoding)
        except UnicodeDecodeError:
            continue
    return raw.decode(encoding="utf-8", errors="ignore")


def _read_text_file(path: Path) -> str:
    for encoding in COMMON_TEXT_ENCODINGS:
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    return path.read_text(encoding="utf-8", errors="ignore")


def _read_json_file(path: Path) -> str:
    raw_text = _read_text_file(path)
    try:
        obj = json.loads(raw_text)
    except json.JSONDecodeError:
        return raw_text
    return json.dumps(obj, ensure_ascii=False, indent=2)


def _extract_pdf_text_pages(path: Path) -> List[str]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise ImportError("Missing dependency `pypdf` for PDF parsing.") from exc

    reader = PdfReader(str(path))
    pages: List[str] = []
    for page in reader.pages:
        pages.append((page.extract_text() or "").strip())
    return pages


def _ocr_pdf_pages_tesseract(path: Path, options: PDFParseOptions) -> List[str]:
    try:
        import pytesseract
    except ImportError as exc:
        raise ImportError("Missing dependency `pytesseract` for OCR fallback.") from exc

    try:
        import pypdfium2 as pdfium
    except ImportError as exc:
        raise ImportError("Missing dependency `pypdfium2` for PDF render before OCR.") from exc

    if options.ocr_tesseract_cmd:
        pytesseract.pytesseract.tesseract_cmd = options.ocr_tesseract_cmd

    try:
        _ = pytesseract.get_tesseract_version()
    except Exception as exc:  # noqa: BLE001
        raise RuntimeError(
            "Tesseract OCR executable not found or unavailable. "
            "Please install Tesseract and set PDF_OCR_TESSERACT_CMD in .env if needed."
        ) from exc

    doc = pdfium.PdfDocument(str(path))
    total_pages = len(doc)
    max_pages = options.ocr_max_pages if options.ocr_max_pages > 0 else total_pages
    pages_to_ocr = min(total_pages, max_pages)

    texts: List[str] = []
    scale = max(72, options.ocr_dpi) / 72.0
    for i in range(pages_to_ocr):
        page = doc[i]
        bitmap = page.render(scale=scale)
        image = bitmap.to_pil()
        text = pytesseract.image_to_string(image, lang=options.ocr_lang).strip()
        texts.append(text)
        try:
            image.close()
        except Exception:
            pass
    return texts


def _read_pdf_file(path: Path, options: PDFParseOptions | None = None) -> str:
    opts = options or PDFParseOptions()
    page_texts = _extract_pdf_text_pages(path)

    extracted_text = "\n\n".join(
        f"[Page {idx}]\n{text}" for idx, text in enumerate(page_texts, start=1) if text
    ).strip()
    if len(extracted_text) >= opts.text_min_chars:
        return extracted_text

    if not opts.ocr_fallback_enabled:
        return extracted_text

    if opts.ocr_engine != "tesseract":
        raise ValueError(f"Unsupported PDF OCR engine: {opts.ocr_engine}")

    ocr_page_texts = _ocr_pdf_pages_tesseract(path, options=opts)
    ocr_text = "\n\n".join(
        f"[Page {idx} OCR]\n{text}" for idx, text in enumerate(ocr_page_texts, start=1) if text
    ).strip()
    return ocr_text or extracted_text


def _read_docx_file(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise ImportError("Missing dependency `python-docx` for DOCX parsing.") from exc

    doc = Document(str(path))
    parts: List[str] = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table_idx, table in enumerate(doc.tables, start=1):
        rows: List[str] = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            values = [value for value in cells if value]
            if values:
                rows.append(" | ".join(values))
        if rows:
            parts.append(f"[Table {table_idx}]\n" + "\n".join(rows))

    return "\n\n".join(parts)


def _read_doc_via_command(path: Path, command: List[str], tool_label: str) -> str:
    completed = subprocess.run(
        command,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        check=False,
    )
    if completed.returncode != 0:
        stderr = _decode_text_bytes(completed.stderr).strip()
        raise RuntimeError(f"{tool_label} failed with code {completed.returncode}: {stderr}")

    text = _decode_text_bytes(completed.stdout).strip()
    if text:
        return text

    stderr = _decode_text_bytes(completed.stderr).strip()
    if stderr:
        raise RuntimeError(f"{tool_label} produced empty output: {stderr}")
    raise RuntimeError(f"{tool_label} produced empty output.")


def _read_doc_via_word_com(path: Path) -> str:
    try:
        import pythoncom
        from win32com import client as win32_client
    except ImportError as exc:
        raise RuntimeError("pywin32 is not installed for Word COM parsing.") from exc

    with tempfile.TemporaryDirectory(prefix="rag_doc_") as temp_dir:
        temp_path = Path(temp_dir)
        output_txt = temp_path / f"{path.stem}.txt"
        word_app = None
        document = None
        pythoncom.CoInitialize()
        try:
            word_app = win32_client.DispatchEx("Word.Application")
            word_app.Visible = False
            word_app.DisplayAlerts = 0
            document = word_app.Documents.Open(str(path.resolve()), ReadOnly=True)
            # wdFormatText = 2
            document.SaveAs(str(output_txt), FileFormat=2)
        finally:
            if document is not None:
                document.Close(False)
            if word_app is not None:
                word_app.Quit()
            pythoncom.CoUninitialize()

        if not output_txt.exists():
            raise RuntimeError("Word COM conversion completed but no output text file was generated.")
        return _read_text_file(output_txt)


def _read_doc_via_antiword(path: Path) -> str:
    antiword_bin = os.getenv("DOC_PARSER_ANTIWORD_PATH") or shutil.which("antiword")
    if not antiword_bin and sys.platform.startswith("win"):
        for candidate in WINDOWS_COMMON_ANTIWORD_PATHS:
            if Path(candidate).exists():
                antiword_bin = candidate
                break
    if not antiword_bin:
        raise RuntimeError("antiword executable not found.")
    return _read_doc_via_command(path, [antiword_bin, str(path.resolve())], tool_label="antiword")


def _read_doc_via_catdoc(path: Path) -> str:
    catdoc_bin = os.getenv("DOC_PARSER_CATDOC_PATH") or shutil.which("catdoc")
    if not catdoc_bin and sys.platform.startswith("win"):
        for candidate in WINDOWS_COMMON_CATDOC_PATHS:
            if Path(candidate).exists():
                catdoc_bin = candidate
                break
    if not catdoc_bin:
        raise RuntimeError("catdoc executable not found.")
    return _read_doc_via_command(path, [catdoc_bin, str(path.resolve())], tool_label="catdoc")


def _read_doc_via_soffice(path: Path) -> str:
    soffice_bin = (
        os.getenv("DOC_PARSER_SOFFICE_PATH")
        or shutil.which("soffice")
        or shutil.which("libreoffice")
    )
    if not soffice_bin and sys.platform.startswith("win"):
        for candidate in WINDOWS_COMMON_SOFFICE_PATHS:
            if Path(candidate).exists():
                soffice_bin = candidate
                break
    if not soffice_bin:
        raise RuntimeError("soffice/libreoffice executable not found.")

    with tempfile.TemporaryDirectory(prefix="rag_doc_") as output_dir:
        output_path = Path(output_dir)
        completed = subprocess.run(
            [
                soffice_bin,
                "--headless",
                "--convert-to",
                "txt:Text",
                "--outdir",
                str(output_path),
                str(path.resolve()),
            ],
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            check=False,
        )
        if completed.returncode != 0:
            stderr = _decode_text_bytes(completed.stderr).strip()
            raise RuntimeError(
                f"LibreOffice conversion failed with code {completed.returncode}: {stderr}"
            )

        preferred_txt = output_path / f"{path.stem}.txt"
        if preferred_txt.exists():
            return _read_text_file(preferred_txt)

        candidates = sorted(output_path.glob("*.txt"))
        if not candidates:
            stdout = _decode_text_bytes(completed.stdout).strip()
            stderr = _decode_text_bytes(completed.stderr).strip()
            raise RuntimeError(
                "LibreOffice conversion succeeded but no txt output was found. "
                f"stdout={stdout} stderr={stderr}"
            )
        return _read_text_file(candidates[0])


def _read_doc_file(path: Path) -> str:
    is_windows = sys.platform.startswith("win")
    strategies: List[tuple[str, Callable[[Path], str]]] = []
    if is_windows:
        strategies.append(("Microsoft Word COM", _read_doc_via_word_com))
    strategies.extend(
        [
            ("antiword", _read_doc_via_antiword),
            ("catdoc", _read_doc_via_catdoc),
            ("LibreOffice", _read_doc_via_soffice),
        ]
    )

    errors: List[str] = []
    for name, reader in strategies:
        try:
            text = reader(path).strip()
            if text:
                return text
            errors.append(f"{name}: empty output")
        except Exception as exc:  # noqa: BLE001
            errors.append(f"{name}: {exc}")

    detail = "; ".join(errors) if errors else "no strategy available"
    raise RuntimeError(
        "Unable to parse .doc file. Install one of the supported extractors: "
        "Microsoft Word + pywin32 (Windows), antiword/catdoc, or LibreOffice. "
        f"file={path} details={detail}"
    )


def _read_xlsx_file(path: Path) -> str:
    try:
        from openpyxl import load_workbook
    except ImportError as exc:
        raise ImportError("Missing dependency `openpyxl` for XLSX parsing.") from exc

    workbook = load_workbook(filename=str(path), read_only=True, data_only=True)
    sheets_text: List[str] = []
    for sheet in workbook.worksheets:
        rows: List[str] = []
        for row in sheet.iter_rows(values_only=True):
            values = [str(cell).strip() for cell in row if cell is not None and str(cell).strip()]
            if values:
                rows.append("\t".join(values))
        if rows:
            sheets_text.append(f"[Sheet: {sheet.title}]\n" + "\n".join(rows))
    workbook.close()
    return "\n\n".join(sheets_text)


def _read_by_suffix(path: Path, pdf_options: PDFParseOptions | None = None) -> str:
    suffix = path.suffix.lower()
    if suffix in TEXT_LIKE_SUFFIXES:
        return _read_text_file(path)
    if suffix == ".json":
        return _read_json_file(path)
    if suffix == ".pdf":
        return _read_pdf_file(path, options=pdf_options)
    if suffix == ".doc":
        return _read_doc_file(path)
    if suffix == ".docx":
        return _read_docx_file(path)
    if suffix == ".xlsx":
        return _read_xlsx_file(path)
    raise ValueError(f"Unsupported file type: {suffix}")


def discover_files(input_dir: Path, extensions: Iterable[str]) -> List[Path]:
    allowed = {ext.lower() for ext in extensions}
    files = []
    for path in input_dir.rglob("*"):
        if (
            path.is_file()
            and not path.name.startswith("~$")
            and path.suffix.lower() in allowed
        ):
            files.append(path)
    files.sort()
    return files


def _stable_doc_id(relative_path: str) -> str:
    digest = hashlib.sha1(relative_path.encode("utf-8")).hexdigest()[:16]
    return f"doc-{digest}"


def _sha256_text(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8")).hexdigest()


def load_documents(
    input_dir: Path,
    extensions: Iterable[str],
    pdf_options: PDFParseOptions | None = None,
) -> List[SourceDocument]:
    files = discover_files(input_dir=input_dir, extensions=extensions)
    documents: List[SourceDocument] = []
    load_errors: List[str] = []

    for file_path in files:
        try:
            text = _read_by_suffix(file_path, pdf_options=pdf_options).strip()
        except Exception as exc:  # noqa: BLE001
            load_errors.append(f"{file_path}: {exc}")
            continue

        if not text:
            continue

        rel_path = file_path.relative_to(input_dir).as_posix()
        documents.append(
            SourceDocument(
                doc_id=_stable_doc_id(rel_path),
                path=str(file_path),
                content=text,
                metadata={
                    "filename": file_path.name,
                    "suffix": file_path.suffix.lower(),
                    "relative_path": rel_path,
                    "content_sha256": _sha256_text(text),
                },
            )
        )

    if load_errors:
        lines = "\n".join(load_errors)
        raise RuntimeError(
            "Some files failed to parse during ingestion. "
            "Please install required dependencies or fix source files.\n"
            f"{lines}"
        )

    return documents
